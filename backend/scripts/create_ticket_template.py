"""
Script to create the boarding pass DOCX template
Based on the SKYLINE AIRWAYS design
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def create_boarding_pass_template():
    """Create SKYLINE AIRWAYS boarding pass template"""
    
    # Create document
    doc = Document()
    
    # Set page margins (Letter size: 8.5" x 11")
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Define colors (matching the React component)
    primary_blue = RGBColor(9, 79, 142)  # hsl(214 100% 35%)
    accent_orange = RGBColor(251, 146, 60)  # hsl(25 95% 53%)
    gray_text = RGBColor(107, 114, 128)  # muted foreground
    dark_text = RGBColor(38, 42, 51)  # foreground
    
    # ===== HEADER SECTION =====
    header_para = doc.add_paragraph()
    header_run = header_para.add_run("SKYLINE AIRWAYS")
    header_run.font.size = Pt(32)
    header_run.font.bold = True
    header_run.font.color.rgb = primary_blue
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact info
    contact_para = doc.add_paragraph()
    contact_run = contact_para.add_run("skylineairways.com | 222 555 7777")
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = gray_text
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # ===== PASSENGER INFORMATION SECTION =====
    passenger_header = doc.add_paragraph("Passenger Information")
    passenger_header.runs[0].font.size = Pt(18)
    passenger_header.runs[0].font.bold = True
    doc.add_paragraph()  # Spacer
    
    # Create table for passenger info
    passenger_table = doc.add_table(rows=2, cols=2)
    passenger_table.style = 'Light Grid Accent 1'
    
    # Passenger Name
    name_cell = passenger_table.rows[0].cells[0]
    name_label = name_cell.add_paragraph("Passenger Name")
    name_label.runs[0].font.size = Pt(11)
    name_label.runs[0].font.color.rgb = gray_text
    name_label.runs[0].font.bold = True
    
    name_value = name_cell.add_paragraph("{{PASSENGER_NAME}}")
    name_value.runs[0].font.size = Pt(20)
    name_value.runs[0].font.bold = True
    
    # Ticket Number
    ticket_cell = passenger_table.rows[0].cells[1]
    ticket_label = ticket_cell.add_paragraph("Ticket Number")
    ticket_label.runs[0].font.size = Pt(11)
    ticket_label.runs[0].font.color.rgb = gray_text
    ticket_label.runs[0].font.bold = True
    
    ticket_value = ticket_cell.add_paragraph("Airline Ticket #{{TICKET_NUMBER}}")
    ticket_value.runs[0].font.size = Pt(16)
    ticket_value.runs[0].font.name = 'Courier New'
    
    doc.add_paragraph()  # Spacer
    
    # ===== DEPARTURE SECTION =====
    departure_header = doc.add_paragraph("Departure")
    departure_header.runs[0].font.size = Pt(18)
    departure_header.runs[0].font.bold = True
    departure_header.runs[0].font.color.rgb = primary_blue
    
    # Departure info table
    departure_table = doc.add_table(rows=3, cols=2)
    departure_table.style = 'Light Grid Accent 1'
    
    # Row 1: Airport and Date
    airport_cell = departure_table.rows[0].cells[0]
    airport_label = airport_cell.add_paragraph("Airport")
    airport_label.runs[0].font.size = Pt(11)
    airport_label.runs[0].font.color.rgb = gray_text
    airport_label.runs[0].font.bold = True
    airport_value = airport_cell.add_paragraph("{{DEPARTURE_AIRPORT}}")
    airport_value.runs[0].font.size = Pt(14)
    
    date_cell = departure_table.rows[0].cells[1]
    date_label = date_cell.add_paragraph("Date")
    date_label.runs[0].font.size = Pt(11)
    date_label.runs[0].font.color.rgb = gray_text
    date_label.runs[0].font.bold = True
    date_value = date_cell.add_paragraph("{{DEPARTURE_DATE}}")
    date_value.runs[0].font.size = Pt(14)
    
    # Row 2: Gate and Departure Time
    gate_cell = departure_table.rows[1].cells[0]
    gate_label = gate_cell.add_paragraph("Gate")
    gate_label.runs[0].font.size = Pt(11)
    gate_label.runs[0].font.color.rgb = gray_text
    gate_label.runs[0].font.bold = True
    gate_value = gate_cell.add_paragraph("{{GATE}}")
    gate_value.runs[0].font.size = Pt(24)
    gate_value.runs[0].font.bold = True
    gate_value.runs[0].font.color.rgb = primary_blue
    
    dep_time_cell = departure_table.rows[1].cells[1]
    dep_time_label = dep_time_cell.add_paragraph("Departure Time")
    dep_time_label.runs[0].font.size = Pt(11)
    dep_time_label.runs[0].font.color.rgb = gray_text
    dep_time_label.runs[0].font.bold = True
    dep_time_value = dep_time_cell.add_paragraph("{{DEPARTURE_TIME}}")
    dep_time_value.runs[0].font.size = Pt(14)
    
    # Row 3: Boarding Time
    boarding_cell = departure_table.rows[2].cells[0]
    boarding_label = boarding_cell.add_paragraph("Boarding Time")
    boarding_label.runs[0].font.size = Pt(11)
    boarding_label.runs[0].font.color.rgb = gray_text
    boarding_label.runs[0].font.bold = True
    boarding_value = boarding_cell.add_paragraph("{{BOARDING_TIME}}")
    boarding_value.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()  # Spacer
    
    # ===== ARRIVAL SECTION =====
    arrival_header = doc.add_paragraph("Arrival")
    arrival_header.runs[0].font.size = Pt(18)
    arrival_header.runs[0].font.bold = True
    arrival_header.runs[0].font.color.rgb = accent_orange
    
    # Arrival info table
    arrival_table = doc.add_table(rows=2, cols=2)
    arrival_table.style = 'Light Grid Accent 1'
    
    # Row 1: Airport and Date
    arr_airport_cell = arrival_table.rows[0].cells[0]
    arr_airport_label = arr_airport_cell.add_paragraph("Airport")
    arr_airport_label.runs[0].font.size = Pt(11)
    arr_airport_label.runs[0].font.color.rgb = gray_text
    arr_airport_label.runs[0].font.bold = True
    arr_airport_value = arr_airport_cell.add_paragraph("{{ARRIVAL_AIRPORT}}")
    arr_airport_value.runs[0].font.size = Pt(14)
    
    arr_date_cell = arrival_table.rows[0].cells[1]
    arr_date_label = arr_date_cell.add_paragraph("Date")
    arr_date_label.runs[0].font.size = Pt(11)
    arr_date_label.runs[0].font.color.rgb = gray_text
    arr_date_label.runs[0].font.bold = True
    arr_date_value = arr_date_cell.add_paragraph("{{ARRIVAL_DATE}}")
    arr_date_value.runs[0].font.size = Pt(14)
    
    # Row 2: Expected Arrival Time
    arr_time_cell = arrival_table.rows[1].cells[0]
    arr_time_label = arr_time_cell.add_paragraph("Expected Arrival Time")
    arr_time_label.runs[0].font.size = Pt(11)
    arr_time_label.runs[0].font.color.rgb = gray_text
    arr_time_label.runs[0].font.bold = True
    arr_time_value = arr_time_cell.add_paragraph("{{ARRIVAL_TIME}}")
    arr_time_value.runs[0].font.size = Pt(14)
    
    doc.add_paragraph()  # Spacer
    
    # ===== FLIGHT INFORMATION SECTION =====
    flight_header = doc.add_paragraph("Flight Information")
    flight_header.runs[0].font.size = Pt(18)
    flight_header.runs[0].font.bold = True
    flight_header.runs[0].font.color.rgb = primary_blue
    
    # Flight info table
    flight_table = doc.add_table(rows=2, cols=2)
    flight_table.style = 'Light Grid Accent 1'
    
    # Row 1: Flight Number and Class
    flight_num_cell = flight_table.rows[0].cells[0]
    flight_num_label = flight_num_cell.add_paragraph("Flight Number")
    flight_num_label.runs[0].font.size = Pt(11)
    flight_num_label.runs[0].font.color.rgb = gray_text
    flight_num_label.runs[0].font.bold = True
    flight_num_value = flight_num_cell.add_paragraph("{{FLIGHT_NUMBER}}")
    flight_num_value.runs[0].font.size = Pt(18)
    flight_num_value.runs[0].font.bold = True
    flight_num_value.runs[0].font.color.rgb = primary_blue
    
    class_cell = flight_table.rows[0].cells[1]
    class_label = class_cell.add_paragraph("Class")
    class_label.runs[0].font.size = Pt(11)
    class_label.runs[0].font.color.rgb = gray_text
    class_label.runs[0].font.bold = True
    class_value = class_cell.add_paragraph("{{CLASS}}")
    class_value.runs[0].font.size = Pt(14)
    
    # Row 2: Aircraft Type and Seat
    aircraft_cell = flight_table.rows[1].cells[0]
    aircraft_label = aircraft_cell.add_paragraph("Aircraft Type")
    aircraft_label.runs[0].font.size = Pt(11)
    aircraft_label.runs[0].font.color.rgb = gray_text
    aircraft_label.runs[0].font.bold = True
    aircraft_value = aircraft_cell.add_paragraph("{{AIRCRAFT_TYPE}}")
    aircraft_value.runs[0].font.size = Pt(14)
    
    seat_cell = flight_table.rows[1].cells[1]
    seat_label = seat_cell.add_paragraph("Seat")
    seat_label.runs[0].font.size = Pt(11)
    seat_label.runs[0].font.color.rgb = gray_text
    seat_label.runs[0].font.bold = True
    seat_value = seat_cell.add_paragraph("{{SEAT}}")
    seat_value.runs[0].font.size = Pt(24)
    seat_value.runs[0].font.bold = True
    seat_value.runs[0].font.color.rgb = primary_blue
    
    doc.add_paragraph()  # Spacer
    
    # ===== BOARDING PASS SECTION =====
    # Create a highlighted section for boarding pass
    boarding_section = doc.add_paragraph()
    boarding_section_format = boarding_section.paragraph_format
    boarding_section_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a table with gradient-like background (we'll use shading)
    boarding_table = doc.add_table(rows=2, cols=4)
    
    # Style the table
    for row in boarding_table.rows:
        for cell in row.cells:
            # Set cell shading to primary blue
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), '00638E')  # Primary blue hex
            cell._element.get_or_add_tcPr().append(shading)
    
    # Row 1: Flight, Gate, Seat, Boarding
    flight_bp_cell = boarding_table.rows[0].cells[0]
    flight_bp_label = flight_bp_cell.add_paragraph("Flight")
    flight_bp_label.runs[0].font.size = Pt(10)
    flight_bp_label.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    flight_bp_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flight_bp_value = flight_bp_cell.add_paragraph("{{FLIGHT_NUMBER}}")
    flight_bp_value.runs[0].font.size = Pt(18)
    flight_bp_value.runs[0].font.bold = True
    flight_bp_value.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    flight_bp_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    gate_bp_cell = boarding_table.rows[0].cells[1]
    gate_bp_label = gate_bp_cell.add_paragraph("Gate")
    gate_bp_label.runs[0].font.size = Pt(10)
    gate_bp_label.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    gate_bp_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gate_bp_value = gate_bp_cell.add_paragraph("{{GATE}}")
    gate_bp_value.runs[0].font.size = Pt(18)
    gate_bp_value.runs[0].font.bold = True
    gate_bp_value.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    gate_bp_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    seat_bp_cell = boarding_table.rows[0].cells[2]
    seat_bp_label = seat_bp_cell.add_paragraph("Seat")
    seat_bp_label.runs[0].font.size = Pt(10)
    seat_bp_label.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    seat_bp_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    seat_bp_value = seat_bp_cell.add_paragraph("{{SEAT}}")
    seat_bp_value.runs[0].font.size = Pt(18)
    seat_bp_value.runs[0].font.bold = True
    seat_bp_value.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    seat_bp_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    boarding_bp_cell = boarding_table.rows[0].cells[3]
    boarding_bp_label = boarding_bp_cell.add_paragraph("Boarding")
    boarding_bp_label.runs[0].font.size = Pt(10)
    boarding_bp_label.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    boarding_bp_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    boarding_bp_value = boarding_bp_cell.add_paragraph("{{BOARDING_TIME}}")
    boarding_bp_value.runs[0].font.size = Pt(16)
    boarding_bp_value.runs[0].font.bold = True
    boarding_bp_value.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    boarding_bp_value.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Row 2: BOARDING PASS text (merged cells)
    bp_text_cell = boarding_table.rows[1].cells[0]
    # Merge all cells in row 2
    for i in range(1, 4):
        bp_text_cell.merge(boarding_table.rows[1].cells[i])
    
    bp_text_para = bp_text_cell.add_paragraph("BOARDING PASS")
    bp_text_para.runs[0].font.size = Pt(24)
    bp_text_para.runs[0].font.bold = True
    bp_text_para.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    bp_text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # ===== QR CODE SECTION =====
    qr_section = doc.add_paragraph()
    qr_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    qr_label = doc.add_paragraph("QR Code:")
    qr_label.runs[0].font.size = Pt(10)
    qr_label.runs[0].font.color.rgb = gray_text
    qr_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add placeholder for QR code
    qr_para = doc.add_paragraph("{{QR_CODE_IMAGE}}")
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    template_dir = Path(__file__).parent.parent / "templates" / "tickets"
    template_dir.mkdir(parents=True, exist_ok=True)
    template_path = template_dir / "boarding_pass_template.docx"
    
    doc.save(str(template_path))
    print(f"✅ Template created successfully at: {template_path}")
    print(f"📄 File: {template_path}")
    
    return template_path


if __name__ == "__main__":
    create_boarding_pass_template()

