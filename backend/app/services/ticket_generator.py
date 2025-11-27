"""
DOCX-based Ticket Generator Service
Generates boarding passes from DOCX templates with placeholder replacement
"""
import os
import logging
import re
from pathlib import Path
from io import BytesIO
from datetime import datetime, timedelta
from typing import Dict, Optional

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import qrcode
from PIL import Image

logger = logging.getLogger(__name__)

# Template directory path
TEMPLATE_DIR = Path(__file__).parent.parent.parent / "templates" / "tickets"
DEFAULT_TEMPLATE = "boarding_pass_template.docx"


class TicketGenerator:
    """Generate boarding pass tickets from DOCX templates"""
    
    def __init__(self, template_path: Optional[str] = None):
        """
        Initialize ticket generator with template path
        
        Args:
            template_path: Path to DOCX template file. If None, uses default template.
        """
        if template_path:
            self.template_path = Path(template_path)
        else:
            self.template_path = TEMPLATE_DIR / DEFAULT_TEMPLATE
        
        if not self.template_path.exists():
            raise FileNotFoundError(
                f"Template not found: {self.template_path}. "
                f"Please create the template file or specify a custom path."
            )
    
    def generate_ticket_data(self, booking, flight, passenger, plane) -> Dict[str, str]:
        """
        Prepare ticket data dictionary from booking/flight/passenger data
        
        Args:
            booking: Booking model instance
            flight: Flight model instance
            passenger: User model instance (passenger)
            plane: Plane model instance
        
        Returns:
            Dictionary with all placeholder values
        """
        # Generate ticket number
        ticket_number = f"PVT-{booking.booked_at.year}-{str(booking.id).zfill(6)}"
        
        # Passenger name (use booking details if available, otherwise user details)
        passenger_name = booking.passenger_name or passenger.full_name or passenger.email
        
        # Airport codes (simple extraction)
        def get_airport_code(name: str) -> str:
            """Extract airport code from airport name"""
            airport_map = {
                'mumbai': 'BOM', 'delhi': 'DEL', 'bangalore': 'BLR', 'chennai': 'MAA',
                'kolkata': 'CCU', 'hyderabad': 'HYD', 'pune': 'PNQ', 'goa': 'GOI',
                'los angeles': 'LAX', 'new york': 'JFK', 'london': 'LHR', 'dubai': 'DXB',
                'singapore': 'SIN', 'tokyo': 'NRT', 'paris': 'CDG', 'frankfurt': 'FRA',
                'teterboro': 'TEB'
            }
            name_lower = name.lower()
            for key, code in airport_map.items():
                if key in name_lower:
                    return code
            return name[:3].upper().replace(' ', '')
        
        origin_code = get_airport_code(flight.origin)
        dest_code = get_airport_code(flight.destination)
        
        # Format dates and times
        departure_date = flight.departure_time.strftime('%d %b %Y').upper()
        departure_time = flight.departure_time.strftime('%H:%M')
        arrival_date = flight.arrival_time.strftime('%d %b %Y').upper()
        arrival_time = flight.arrival_time.strftime('%H:%M')
        
        # Flight number
        flight_number = getattr(flight, 'flight_number', None) or f"PJ-{flight.id}"
        
        # Seat assignment
        if booking.is_full_charter:
            seat_number = "CHARTER"
        elif booking.seat_id and booking.seat:
            seat_number = booking.seat.seat_number
        else:
            seat_number = "N/A"
        
        # Aircraft type
        aircraft_type = plane.model if plane else "Private Jet"
        
        # Class (determine from seat or booking)
        if booking.seat_id and booking.seat:
            class_type = booking.seat.class_type.upper() if booking.seat.class_type else "FIRST CLASS"
        else:
            class_type = "FIRST CLASS"  # Default for private jets
        
        # Gate (default or from flight if available)
        gate = getattr(flight, 'gate', None) or "G4"
        
        # Boarding time (typically 30 minutes before departure)
        boarding_datetime = flight.departure_time - timedelta(minutes=30)
        boarding_time = boarding_datetime.strftime('%H:%M')
        
        return {
            "PASSENGER_NAME": passenger_name.upper(),
            "TICKET_NUMBER": ticket_number,
            "DEPARTURE_AIRPORT": flight.origin.upper(),
            "DEPARTURE_AIRPORT_CODE": origin_code,
            "DEPARTURE_DATE": departure_date,
            "DEPARTURE_TIME": departure_time,
            "GATE": gate,
            "BOARDING_TIME": boarding_time,
            "FLIGHT_NUMBER": flight_number,
            "CLASS": class_type,
            "AIRCRAFT_TYPE": aircraft_type,
            "SEAT": seat_number,
            "ARRIVAL_AIRPORT": flight.destination.upper(),
            "ARRIVAL_AIRPORT_CODE": dest_code,
            "ARRIVAL_DATE": arrival_date,
            "ARRIVAL_TIME": arrival_time,
        }
    
    def generate_qr_code(self, data: str, size: int = 200) -> BytesIO:
        """
        Generate QR code image
        
        Args:
            data: Data to encode in QR code
            size: QR code size in pixels
        
        Returns:
            BytesIO buffer containing PNG image
        """
        try:
            qr = qrcode.QRCode(version=1, box_size=10, border=4)
            qr.add_data(data)
            qr.make(fit=True)
            img = qr.make_image(fill_color="black", back_color="white")
            
            # Resize if needed
            if size != 200:
                img = img.resize((size, size), Image.Resampling.LANCZOS)
            
            buffer = BytesIO()
            img.save(buffer, format='PNG')
            buffer.seek(0)
            return buffer
        except Exception as e:
            logger.error(f"Error generating QR code: {e}")
            raise
    
    def replace_placeholders_in_docx(self, doc: Document, data: Dict[str, str], qr_code_buffer: Optional[BytesIO] = None) -> Document:
        """
        Replace placeholders in DOCX document with actual values
        
        Args:
            doc: Document object to modify
            data: Dictionary of placeholder -> value mappings
            qr_code_buffer: Optional QR code image buffer to insert
        
        Returns:
            Modified document
        """
        # Replace text placeholders in paragraphs
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            new_text = original_text
            
            # Replace all placeholders
            for placeholder, value in data.items():
                placeholder_pattern = f"{{{{{placeholder}}}}}"
                if placeholder_pattern in new_text:
                    new_text = new_text.replace(placeholder_pattern, str(value))
            
            if new_text != original_text:
                # Clear paragraph and add new text while preserving formatting
                if paragraph.runs:
                    # Keep first run's formatting
                    first_run = paragraph.runs[0]
                    paragraph.clear()
                    new_run = paragraph.add_run(new_text)
                    # Copy formatting from first run
                    new_run.font.name = first_run.font.name
                    new_run.font.size = first_run.font.size
                    new_run.font.bold = first_run.font.bold
                    new_run.font.italic = first_run.font.italic
                else:
                    paragraph.text = new_text
            
            # Handle QR_CODE_IMAGE placeholder
            if qr_code_buffer and "{{QR_CODE_IMAGE}}" in paragraph.text:
                # Insert QR code image
                paragraph.clear()
                run = paragraph.add_run()
                qr_code_buffer.seek(0)
                run.add_picture(qr_code_buffer, width=Inches(1.0))
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original_text = paragraph.text
                        new_text = original_text
                        
                        # Replace all placeholders
                        for placeholder, value in data.items():
                            placeholder_pattern = f"{{{{{placeholder}}}}}"
                            if placeholder_pattern in new_text:
                                new_text = new_text.replace(placeholder_pattern, str(value))
                        
                        if new_text != original_text:
                            if paragraph.runs:
                                first_run = paragraph.runs[0]
                                paragraph.clear()
                                new_run = paragraph.add_run(new_text)
                                new_run.font.name = first_run.font.name
                                new_run.font.size = first_run.font.size
                                new_run.font.bold = first_run.font.bold
                                new_run.font.italic = first_run.font.italic
                            else:
                                paragraph.text = new_text
                        
                        # Handle QR code in table cells
                        if qr_code_buffer and "{{QR_CODE_IMAGE}}" in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run()
                            qr_code_buffer.seek(0)
                            run.add_picture(qr_code_buffer, width=Inches(1.0))
        
        # Replace placeholders in headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for paragraph in section.header.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text
                    for placeholder, value in data.items():
                        placeholder_pattern = f"{{{{{placeholder}}}}}"
                        new_text = new_text.replace(placeholder_pattern, str(value))
                    if new_text != original_text:
                        paragraph.text = new_text
            
            # Footer
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    original_text = paragraph.text
                    new_text = original_text
                    for placeholder, value in data.items():
                        placeholder_pattern = f"{{{{{placeholder}}}}}"
                        new_text = new_text.replace(placeholder_pattern, str(value))
                    if new_text != original_text:
                        paragraph.text = new_text
        
        return doc
    
    def generate_ticket_docx(self, booking, flight, passenger, plane, qr_code_data: Optional[str] = None) -> BytesIO:
        """
        Generate ticket as DOCX document
        
        Args:
            booking: Booking model instance
            flight: Flight model instance
            passenger: User model instance
            plane: Plane model instance
            qr_code_data: Optional QR code data string. If None, generates automatically.
        
        Returns:
            BytesIO buffer containing DOCX document
        """
        try:
            # Load template
            logger.info(f"Loading template from: {self.template_path}")
            doc = Document(str(self.template_path))
            
            # Prepare ticket data
            ticket_data = self.generate_ticket_data(booking, flight, passenger, plane)
            
            # Generate QR code if data provided
            qr_code_buffer = None
            if qr_code_data or True:  # Always generate QR code
                if not qr_code_data:
                    # Auto-generate QR code data
                    qr_code_data = (
                        f"{ticket_data['TICKET_NUMBER']}|"
                        f"{ticket_data['PASSENGER_NAME']}|"
                        f"{ticket_data['FLIGHT_NUMBER']}|"
                        f"{ticket_data['DEPARTURE_AIRPORT_CODE']}|"
                        f"{ticket_data['ARRIVAL_AIRPORT_CODE']}|"
                        f"{ticket_data['DEPARTURE_DATE']}|"
                        f"{ticket_data['DEPARTURE_TIME']}"
                    )
                qr_code_buffer = self.generate_qr_code(qr_code_data)
            
            # Replace placeholders
            doc = self.replace_placeholders_in_docx(doc, ticket_data, qr_code_buffer)
            
            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            logger.info("DOCX ticket generated successfully")
            return buffer
            
        except Exception as e:
            logger.error(f"Error generating DOCX ticket: {e}")
            raise
    
    def docx_to_pdf(self, docx_buffer: BytesIO) -> BytesIO:
        """
        Convert DOCX to PDF
        
        This method tries multiple conversion approaches:
        1. LibreOffice headless (if available)
        2. Falls back to returning DOCX if PDF conversion fails
        
        Args:
            docx_buffer: BytesIO buffer containing DOCX document
        
        Returns:
            BytesIO buffer containing PDF document
        """
        import subprocess
        import tempfile
        import shutil
        
        try:
            # Try LibreOffice headless conversion
            libreoffice_path = self._find_libreoffice()
            if libreoffice_path:
                logger.info("Using LibreOffice for DOCX to PDF conversion")
                
                # Create temporary files
                with tempfile.TemporaryDirectory() as temp_dir:
                    # Save DOCX to temp file
                    docx_path = os.path.join(temp_dir, "ticket.docx")
                    pdf_path = os.path.join(temp_dir, "ticket.pdf")
                    
                    docx_buffer.seek(0)
                    with open(docx_path, 'wb') as f:
                        f.write(docx_buffer.read())
                    
                    # Convert using LibreOffice
                    cmd = [
                        libreoffice_path,
                        "--headless",
                        "--convert-to", "pdf",
                        "--outdir", temp_dir,
                        docx_path
                    ]
                    
                    result = subprocess.run(
                        cmd,
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    if result.returncode == 0 and os.path.exists(pdf_path):
                        # Read PDF into buffer
                        pdf_buffer = BytesIO()
                        with open(pdf_path, 'rb') as f:
                            pdf_buffer.write(f.read())
                        pdf_buffer.seek(0)
                        logger.info("PDF conversion successful")
                        return pdf_buffer
                    else:
                        logger.warning(f"LibreOffice conversion failed: {result.stderr}")
            
            # If LibreOffice not available or conversion failed, try alternative methods
            logger.warning("LibreOffice not available. Trying alternative conversion methods...")
            
            # Option 2: Use docx2pdf library (Windows only with win32com)
            try:
                import docx2pdf
                with tempfile.TemporaryDirectory() as temp_dir:
                    docx_path = os.path.join(temp_dir, "ticket.docx")
                    pdf_path = os.path.join(temp_dir, "ticket.pdf")
                    
                    docx_buffer.seek(0)
                    with open(docx_path, 'wb') as f:
                        f.write(docx_buffer.read())
                    
                    docx2pdf.convert(docx_path, pdf_path)
                    
                    if os.path.exists(pdf_path):
                        pdf_buffer = BytesIO()
                        with open(pdf_path, 'rb') as f:
                            pdf_buffer.write(f.read())
                        pdf_buffer.seek(0)
                        logger.info("PDF conversion successful using docx2pdf")
                        return pdf_buffer
            except ImportError:
                logger.warning("docx2pdf library not available")
            except Exception as e:
                logger.warning(f"docx2pdf conversion failed: {e}")
            
            # If all methods fail, return DOCX and log error
            logger.error("All PDF conversion methods failed. Returning DOCX instead.")
            docx_buffer.seek(0)
            return docx_buffer
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            # Return DOCX as fallback
            docx_buffer.seek(0)
            return docx_buffer
    
    def _find_libreoffice(self) -> Optional[str]:
        """Find LibreOffice executable path"""
        import shutil
        
        # Common paths
        paths = [
            "soffice",  # In PATH
            "/usr/bin/soffice",  # Linux
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",  # Windows
            "C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe",  # Windows 32-bit
        ]
        
        for path in paths:
            if shutil.which(path) or (os.path.exists(path) and os.path.isfile(path)):
                return path
        
        return None
    
    def generate_ticket_pdf(self, booking, flight, passenger, plane, qr_code_data: Optional[str] = None, prefer_pdf: bool = True) -> tuple[BytesIO, str]:
        """
        Generate ticket and convert to PDF if possible
        
        Args:
            booking: Booking model instance
            flight: Flight model instance
            passenger: User model instance
            plane: Plane model instance
            qr_code_data: Optional QR code data string
            prefer_pdf: If True, tries to convert to PDF. If False or conversion fails, returns DOCX.
        
        Returns:
            Tuple of (BytesIO buffer, media_type)
        """
        # Generate DOCX
        docx_buffer = self.generate_ticket_docx(booking, flight, passenger, plane, qr_code_data)
        
        # Try to convert to PDF if preferred
        if prefer_pdf:
            pdf_buffer = self.docx_to_pdf(docx_buffer)
            
            # Check if conversion was successful (PDF files start with %PDF)
            pdf_buffer.seek(0)
            first_bytes = pdf_buffer.read(4)
            pdf_buffer.seek(0)
            
            if first_bytes == b'%PDF':
                return pdf_buffer, "application/pdf"
            else:
                # Conversion failed, return DOCX
                docx_buffer.seek(0)
                return docx_buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            return docx_buffer, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

